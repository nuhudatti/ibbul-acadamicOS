# -*- coding: utf-8 -*-
"""Repair and complete Chapter 3 diagrams in the project report."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / 'Nuhu_Muhammad_Datti_Final_Year_Project_Report_FINAL.docx'
SRC = ROOT / 'Nuhu_Muhammad_Datti_Final_Year_Project_Report.docx'
DOC = OUT if OUT.exists() else SRC
DIAG = ROOT / 'report_assets' / 'diagrams'

CH3_DIAGRAMS = [
    ('figure_3_2_use_case.png', 'Figure 3.2: Use case diagram showing actors and major system functions',
     'Figure 3.2 shows the primary actors and their interactions with the system. Students view approved results and participate in e-learning. Examiners manage offerings and grades. HODs upload and approve results. Faculty and Super Administrators manage structure and governance.'),
    ('figure_3_3_erd.png', 'Figure 3.3: Entity Relationship Diagram (ERD) of core database entities',
     'Figure 3.3 illustrates the relational structure of the database. The User entity links to Result, Enrollment, and upload records. Course connects results to LMSOfferings. Foreign keys enforce referential integrity across the four Django applications.'),
    ('figure_3_4_class.png', 'Figure 3.4: UML class diagram of major model classes and relationships',
     'Figure 3.4 models the principal Django classes. User supports role-based authentication. Result, Course, and ResultUploadBatch implement the results module. LMSOffering, Module, Lesson, Quiz, Assignment, and Enrollment implement the learning module.'),
    ('figure_3_5_sequence_result_check.png', 'Figure 3.5: Sequence diagram for student result checking workflow',
     'Figure 3.5 traces the student result checking process from the /results page through the API proxy to ResultViewSet.my_results. Only APPROVED or LOCKED_PUBLISHED results are returned to students.'),
    ('figure_3_6_activity_hod_upload.png', 'Figure 3.6: Activity diagram for HOD result upload, validation, and approval workflow',
     'Figure 3.6 describes the HOD workflow: validate file, review errors, submit valid rows as HOD_REVIEW, approve to LOCKED_PUBLISHED, then students view results.'),
]

LIST_OF_FIGURES = [
    'Figure 3.1: Overall system architecture of the IBBUL Unified Academic Operating System',
    'Figure 3.2: Use case diagram showing actors and major system functions',
    'Figure 3.3: Entity Relationship Diagram (ERD) of core database entities',
    'Figure 3.4: UML class diagram of major model classes and relationships',
    'Figure 3.5: Sequence diagram for student result checking workflow',
    'Figure 3.6: Activity diagram for HOD result upload, validation, and approval workflow',
    'Figure 3.7: Deployment architecture on Vercel, Render, PostgreSQL, and Cloudinary',
]

LIST_OF_TABLES = [
    'Table 1.1: Scope of the Implemented System',
    'Table 1.2: Definition of Terms',
    'Table 2.1: Comparison of Existing Systems and Proposed System',
    'Table 3.1: Major Database Tables',
    'Table 4.1: Implementation Tools',
    'Table 4.2: Functional Testing',
    'Table 4.3: Authentication and Authorization Testing',
    'Table 4.4: Upload Validation Testing',
    'Table C.1: Sample Result Upload Template',
    'Table E.1: Selected API Endpoints',
]


def set_text(p, text):
    for r in p.runs:
        r.text = ''
    (p.runs[0] if p.runs else p.add_run()).text = text


def has_image(p):
    for run in p.runs:
        if run._element.xpath('.//a:blip'):
            return True
        if run._element.xpath('.//pic:pic'):
            return True
    return False


def insert_image_in_paragraph(p, img_path, width=6.2):
    set_text(p, '')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Inches(width))


def add_paragraph_after(ref, text, center=False, bold=False, justify=False):
    new_p = OxmlElement('w:p')
    ref._element.addnext(new_p)
    para = Paragraph(new_p, ref._parent)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    if bold:
        run.bold = True
        run.font.size = Pt(11)
    return para


def in_chapter_three(idx, paras):
    ch3 = ch4 = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == 'CHAPTER THREE':
            ch3 = i
        if t == 'CHAPTER FOUR' and ch3 is not None:
            ch4 = i
            break
    return ch3 is not None and ch4 is not None and ch3 <= idx < ch4


def repair():
    doc = Document(str(DOC))
    paras = list(doc.paragraphs)

    # Fix 3.8 database intro
    for i, p in enumerate(paras):
        if p.text.strip() == '3.8 Database Design' and i + 1 < len(paras):
            nxt = paras[i + 1]
            if 'Figure 3.3' not in nxt.text:
                set_text(nxt, (
                    'The database was designed using normalization principles to reduce duplication and '
                    'maintain consistency. Related records are linked using foreign keys. Figure 3.3 presents '
                    'the entity relationship diagram of the core tables implemented in PostgreSQL. Table 3.1 '
                    'summarises the major tables and their purposes.'
                ))
            break

    # Remove duplicate architecture explanation (keep one)
    seen_arch_exp = False
    for p in paras:
        if p.text.strip().startswith('The architecture is illustrated in Figure 3.1'):
            if seen_arch_exp:
                set_text(p, '')
            seen_arch_exp = True

    # Insert missing diagrams at [See Chapter Three] placeholders in order
    diagram_idx = 0
    for i, p in enumerate(paras):
        if not in_chapter_three(i, paras):
            continue
        if p.text.strip() == '[See Chapter Three for system design diagrams.]':
            if diagram_idx >= len(CH3_DIAGRAMS):
                set_text(p, '')
                continue
            fname, caption, explanation = CH3_DIAGRAMS[diagram_idx]
            img = DIAG / fname
            if img.exists():
                insert_image_in_paragraph(p, img)
                cap = add_paragraph_after(p, caption, center=True, bold=True)
                add_paragraph_after(cap, explanation, justify=True)
                print(f'Inserted {fname} at paragraph {i}')
            diagram_idx += 1

    # Ensure 3.1 has image once in Chapter Three (before first caption)
    arch_done = False
    for i, p in enumerate(paras):
        if not in_chapter_three(i, paras):
            continue
        if p.text.strip().startswith('Figure 3.1: Overall system') and not arch_done:
            prev = paras[i - 1] if i > 0 else None
            if prev and not has_image(prev):
                img = DIAG / 'figure_3_1_system_architecture.png'
                if img.exists():
                    insert_image_in_paragraph(prev, img)
                    print('Inserted architecture image before caption')
            arch_done = True

    # Ensure 3.7 deployment image in section 3.10
    dep_done = False
    for i, p in enumerate(paras):
        if not in_chapter_three(i, paras):
            continue
        if p.text.strip().startswith('Figure 3.7: Deployment') and not dep_done:
            prev = paras[i - 1] if i > 0 else None
            if prev and not has_image(prev):
                img = DIAG / 'figure_3_7_deployment.png'
                if img.exists():
                    insert_image_in_paragraph(prev, img)
                    print('Inserted deployment image before caption')
            dep_done = True

    # Fix LIST OF FIGURES — replace block after heading
    for i, p in enumerate(paras):
        if p.text.strip() != 'LIST OF FIGURES':
            continue
        j = i + 1
        # Remove following paragraphs until CHAPTER ONE
        while j < len(paras) and paras[j].text.strip() not in ('CHAPTER ONE', 'CHAPTER 1'):
            if 'Figure 3.' in paras[j].text or 'Generate automatically' in paras[j].text or paras[j].text.strip() == '':
                set_text(paras[j], '')
            j += 1
        # Insert clean list before CHAPTER ONE
        anchor = None
        for k in range(i + 1, len(paras)):
            if paras[k].text.strip() == 'CHAPTER ONE':
                anchor = paras[k]
                break
        if anchor:
            prev = anchor
            for fig in reversed(LIST_OF_FIGURES):
                new_p = OxmlElement('w:p')
                prev._element.addprevious(new_p)
                fp = Paragraph(new_p, prev._parent)
                fp.add_run(fig)
                prev = fp
        break

    # Appendix diagram placeholders
    in_appendix = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith('APPENDIX'):
            in_appendix = True
        if in_appendix and t.startswith('[Insert') and 'Diagram' in t:
            set_text(p, '(Refer to Chapter Three for the official system design diagrams.)')

    # Fix LIST OF TABLES
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() != 'LIST OF TABLES':
            continue
        anchor = None
        for k in range(i + 1, len(doc.paragraphs)):
            if doc.paragraphs[k].text.strip() == 'LIST OF FIGURES':
                anchor = doc.paragraphs[k]
                break
        if anchor:
            prev = anchor
            for tbl in reversed(LIST_OF_TABLES):
                new_p = OxmlElement('w:p')
                prev._element.addprevious(new_p)
                tp = Paragraph(new_p, prev._parent)
                tp.add_run(tbl)
                prev = tp
        break

    for p in doc.paragraphs:
        if '[MONTH, 2026]' in p.text:
            set_text(p, p.text.replace('[MONTH, 2026]', 'June, 2026'))

    out_path = OUT
    doc.save(str(OUT))
    print(f'Repaired and saved {out_path}')


if __name__ == '__main__':
    repair()
