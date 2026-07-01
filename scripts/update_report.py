# -*- coding: utf-8 -*-
"""
Finalize Nuhu Muhammad Datti Final Year Project Report:
- Insert Chapter 3 UML diagrams with captions
- Fix placeholders, cross-references, list of figures
- Preserve existing writing style
"""
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = ROOT / 'Nuhu_Muhammad_Datti_Final_Year_Project_Report.docx'
DIAG = ROOT / 'report_assets' / 'diagrams'

FIGURES = [
    ('[Insert System Architecture Diagram]', 'figure_3_1_system_architecture.png',
     'Figure 3.1: Overall system architecture of the IBBUL Unified Academic Operating System'),
    ('[Insert Use Case Diagram]', 'figure_3_2_use_case.png',
     'Figure 3.2: Use case diagram showing actors and major system functions'),
    ('[Insert ER Diagram Here]', 'figure_3_3_erd.png',
     'Figure 3.3: Entity Relationship Diagram (ERD) of core database entities'),
    ('[Insert Class Diagram]', 'figure_3_4_class.png',
     'Figure 3.4: UML class diagram of major model classes and relationships'),
    ('[Insert Sequence Diagram]', 'figure_3_5_sequence_result_check.png',
     'Figure 3.5: Sequence diagram for student result checking workflow'),
    ('[Insert Activity Diagram]', 'figure_3_6_activity_hod_upload.png',
     'Figure 3.6: Activity diagram for HOD result upload, validation, and approval workflow'),
    ('[Insert Deployment Architecture Diagram]', 'figure_3_7_deployment.png',
     'Figure 3.7: Deployment architecture on Vercel, Render, PostgreSQL, and Cloudinary'),
]

FIGURE_EXPLANATIONS = {
    '[Insert System Architecture Diagram]': (
        'Figure 3.1 presents the three-tier client–server architecture implemented in this project. '
        'The browser communicates with the Next.js frontend hosted on Vercel. API requests are routed '
        'through the same-origin proxy to the Django REST Framework backend on Render. The backend '
        'persists academic data in PostgreSQL, stores learning media on Cloudinary, and sends '
        'transactional email through SendGrid. Optional Redis supports Celery background tasks for '
        'large uploads and exports.'
    ),
    '[Insert Use Case Diagram]': (
        'Figure 3.2 shows the primary actors and their interactions with the system. Students view '
        'approved results and participate in e-learning activities. Examiners manage course offerings, '
        'lessons, assignments, and quizzes. HODs upload and validate departmental results and approve '
        'them for publication. Faculty and Super Administrators manage academic structure, user '
        'assignments, governance, and audit oversight.'
    ),
    '[Insert ER Diagram Here]': (
        'Figure 3.3 illustrates the relational structure of the database. The User entity is central '
        'to both results and learning modules. Results link students to courses per session and semester. '
        'LMSOffering connects courses to instructors and contains modules and lessons. Enrollments and '
        'StudentCourseRegistration record official and LMS participation respectively. ResultUploadBatch '
        'tracks bulk uploads from HOD users.'
    ),
    '[Insert Class Diagram]': (
        'Figure 3.4 models the main Django classes implemented in the system. The User class supports '
        'role-based authentication. Result, Course, and ResultUploadBatch implement the results module. '
        'LMSOffering, Module, Lesson, Quiz, Assignment, and Enrollment implement the learning module. '
        'Cardinality labels indicate one-to-many and one-to-one relationships as defined in the models.'
    ),
    '[Insert Sequence Diagram]': (
        'Figure 3.5 traces the student result checking process from the /results page through the Axios '
        'client, Vercel API proxy, and Django ResultViewSet.my_results endpoint. The backend returns '
        'only results with status APPROVED or LOCKED_PUBLISHED, together with SemesterSummary records '
        'for GPA display. This ensures students cannot view unapproved HOD_REVIEW records.'
    ),
    '[Insert Activity Diagram]': (
        'Figure 3.6 describes the HOD result upload workflow. The HOD selects a spreadsheet and academic '
        'period, then calls the validate endpoint. The ResultUploadService parses IBBUL-format files and '
        'reports invalid rows. On successful validation, submit creates a ResultUploadBatch and Result '
        'records in HOD_REVIEW status. After HOD approval, status changes to LOCKED_PUBLISHED and '
        'students can access results through the results module.'
    ),
    '[Insert Deployment Architecture Diagram]': (
        'Figure 3.7 shows the production deployment configuration. The Next.js application is hosted on '
        'Vercel for global CDN delivery. The Django API runs on Render with Gunicorn. PostgreSQL is '
        'provided through DATABASE_URL (e.g., Neon). Cloudinary stores lesson videos, PDFs, and platform '
        'branding assets. Redis and Celery are optional for asynchronous batch processing.'
    ),
}

LIST_OF_FIGURES = [
    'Figure 3.1: Overall system architecture of the IBBUL Unified Academic Operating System',
    'Figure 3.2: Use case diagram showing actors and major system functions',
    'Figure 3.3: Entity Relationship Diagram (ERD) of core database entities',
    'Figure 3.4: UML class diagram of major model classes and relationships',
    'Figure 3.5: Sequence diagram for student result checking workflow',
    'Figure 3.6: Activity diagram for HOD result upload, validation, and approval workflow',
    'Figure 3.7: Deployment architecture on Vercel, Render, PostgreSQL, and Cloudinary',
]

LIST_OF_TABLES = [
    'Table 1.1: Scope of the Implemented System',
    'Table 1.2: Definition of Terms',
    'Table 2.1: Comparison of Existing Systems',
    'Table 3.1: Major Database Tables',
    'Table 4.1: Implementation Tools',
    'Table 4.2: Functional Testing',
    'Table 4.3: Authentication and Authorization Testing',
    'Table 4.4: Upload Validation Testing',
    'Table C.1: Sample Result Upload Template',
    'Table E.1: Selected API Endpoints',
]

REPLACEMENTS = {
    '[MATRICULATION NUMBER]': '[MATRICULATION NUMBER]',  # Student to insert official matric number before printing
    '[MONTH, 2026]': 'June, 2026',
    '[Generate automatically in Microsoft Word or Google Docs after opening the document: Insert → Table of contents.]':
        'Right-click this heading in Microsoft Word and select Update Field to refresh the table of contents after final printing.',
    '[Generate automatically after final formatting.]': '',  # handled separately for tables/figures lists
    '[Generate automatically after inserting screenshots and diagrams.]': '',  # handled separately
    'The system architecture diagram should show the browser client, Next.js frontend, Django REST backend, PostgreSQL database, Cloudinary storage, Vercel':
        'The architecture is illustrated in Figure 3.1. It comprises the browser client, Next.js frontend on Vercel, Django REST backend on Render, PostgreSQL database, Cloudinary media storage, and SendGrid email service.',
    'The use case diagram should show the interaction between users and the system. It should include Super Administrator, Faculty Administrator, HOD, Lect':
        'The use case diagram in Figure 3.2 shows interactions between five actor types and the major functional areas of the platform, including result management, e-learning, and governance.',
    'The class diagram should show the major system entities and their relationships. It helps explain the structure of the implemented system.':
        'The UML class diagram in Figure 3.4 represents the principal Django model classes in the accounts, academics, and learning applications.',
    'The sequence diagram should show how a user action, such as result upload, moves from frontend to backend, validation service, database, and response.':
        'Figure 3.5 presents the sequence of messages when a student checks results, from the Next.js results page to the PostgreSQL database.',
    'The activity diagram should show the workflow for login, result upload, validation, submission, and approval.':
        'Figure 3.6 models the HOD result upload workflow, including validation, submission, review, and student publication stages.',
}

# Appendix duplicate placeholders to remove (keep Chapter 3 versions only)
APPENDIX_REMOVE = {
    '[Insert Architecture Diagram]',
    '[Insert ER Diagram Here]',
    '[Insert Use Case Diagram]',
    '[Insert Activity Diagram]',
    '[Insert Class Diagram]',
    '[Insert Sequence Diagram]',
}


def set_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_image_after(paragraph, image_path, width_inches=6.0):
    """Replace placeholder paragraph with centered image."""
    set_paragraph_text(paragraph, '')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def insert_caption_after(paragraph, caption_text, explanation=None):
    """Add caption and optional explanation paragraph after image paragraph."""
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    parent = paragraph._element.getparent()
    doc = paragraph.part.document if hasattr(paragraph.part, 'document') else None

    # Caption paragraph
    cap_p = OxmlElement('w:p')
    paragraph._element.addnext(cap_p)
    cap = Paragraph(cap_p, paragraph._parent)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(11)

    if explanation:
        exp_p = OxmlElement('w:p')
        cap._element.addnext(exp_p)
        exp = Paragraph(exp_p, paragraph._parent)
        exp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        exp.add_run(explanation)


def fill_list_section(doc, heading_text, items, start_marker):
    """Replace placeholder under LIST OF TABLES/FIGURES."""
    found = False
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text:
            found = True
            continue
        if found and start_marker in p.text:
            set_paragraph_text(p, items[0] if items else '')
            prev = p
            for item in items[1:]:
                new_p = deepcopy(p._element)
                prev._element.addnext(new_p)
                # get the new paragraph object
                from docx.text.paragraph import Paragraph
                np = Paragraph(new_p, doc)
                np.add_run(item)
                prev = np
            return


def process_document():
    doc = Document(str(DOC_PATH))

    # Global text replacements
    for p in doc.paragraphs:
        t = p.text.strip()
        for old, new in REPLACEMENTS.items():
            if old in p.text:
                set_paragraph_text(p, p.text.replace(old, new))
        # Remove appendix duplicate diagram placeholders
        if t in APPENDIX_REMOVE:
            set_paragraph_text(p, '[See Chapter Three for system design diagrams.]')

    # Insert diagrams at Chapter 3 placeholders
    for placeholder, filename, caption in FIGURES:
        img = DIAG / filename
        if not img.exists():
            print(f'WARNING: missing {img}')
            continue
        for p in doc.paragraphs:
            if p.text.strip() == placeholder:
                insert_image_after(p, img, 6.2)
                insert_caption_after(p, caption, FIGURE_EXPLANATIONS.get(placeholder))
                print(f'Inserted {filename}')
                break

    # Enhance section 3.9 headings
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == '3.9 UML Diagrams':
            # add subsections after chapter summary is far - we'll update text inline
            pass

    # Update 3.6, 3.7, 3.8, 3.10 intro paragraphs with figure references
    section_updates = {
        '3.6 System Architecture': (
            'The system follows a client–server architecture as shown in Figure 3.1. The frontend provides '
            'the user interface and sends HTTP requests to the backend through a secure proxy. The backend '
            'processes business logic, enforces role-based access control, and communicates with PostgreSQL '
            'and Cloudinary. JSON Web Tokens are used for stateless authentication between the client and API.'
        ),
        '3.8 Database Design': (
            'The database was designed using normalization principles to reduce duplication and maintain '
            'consistency. Related records are linked using foreign keys. Figure 3.3 presents the entity '
            'relationship diagram of the core tables implemented in PostgreSQL. Table 3.1 summarises the '
            'major tables and their purposes.'
        ),
        '3.10 Deployment Architecture': (
            'The frontend is deployed on Vercel, while the backend is deployed on Render. PostgreSQL is '
            'used as the production database through DATABASE_URL, and Cloudinary is used for media storage. '
            'Figure 3.7 illustrates the deployment tiers and communication paths between components.'
        ),
    }
    paras = list(doc.paragraphs)
    for idx, p in enumerate(paras):
        style = p.style.name if p.style else ''
        heading = p.text.strip()
        if style == 'Heading 2' and heading in section_updates:
            if idx + 1 < len(paras):
                nxt = paras[idx + 1]
                if nxt.style.name == 'Normal' and not nxt.text.strip().startswith('Figure'):
                    set_paragraph_text(nxt, section_updates[heading])

    # List of figures
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'LIST OF FIGURES':
            if i + 1 < len(doc.paragraphs):
                fp = doc.paragraphs[i + 1]
                if 'Generate automatically' in fp.text or fp.text.strip() == '':
                    set_paragraph_text(fp, LIST_OF_FIGURES[0])
                    prev = fp
                    for fig in LIST_OF_FIGURES[1:]:
                        new_el = deepcopy(fp._element)
                        prev._element.addnext(new_el)
                        from docx.text.paragraph import Paragraph
                        np = Paragraph(new_el, doc)
                        np.add_run(fig)
                        prev = np
            break

    # List of tables
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'LIST OF TABLES':
            if i + 1 < len(doc.paragraphs):
                tp = doc.paragraphs[i + 1]
                if 'Generate automatically' in tp.text or tp.text.strip() == '':
                    set_paragraph_text(tp, LIST_OF_TABLES[0])
                    prev = tp
                    for tbl in LIST_OF_TABLES[1:]:
                        new_el = deepcopy(tp._element)
                        prev._element.addnext(new_el)
                        from docx.text.paragraph import Paragraph
                        np = Paragraph(new_el, doc)
                        np.add_run(tbl)
                        prev = np
            break

    # Update Chapter 3 summary
    for p in doc.paragraphs:
        if p.text.strip().startswith('This chapter discussed system analysis and design'):
            set_paragraph_text(p, (
                'This chapter discussed system analysis and design. It explained the methodology, existing '
                'system, proposed system, architecture (Figure 3.1), use cases (Figure 3.2), database design '
                '(Figure 3.3 and Table 3.1), UML diagrams (Figures 3.4–3.6), and deployment architecture '
                '(Figure 3.7). These design artefacts guided the implementation presented in Chapter Four.'
            ))

    # Update 1.10 organization to mention diagrams
    for p in doc.paragraphs:
        if p.text.strip().startswith('Chapter One introduces the study'):
            set_paragraph_text(p, (
                'Chapter One introduces the study. Chapter Two reviews related literature and existing systems. '
                'Chapter Three explains system analysis and design, including architecture, database design, and '
                'UML diagrams. Chapter Four discusses system implementation and testing. Chapter Five presents '
                'the summary, conclusion, and recommendations.'
            ))

    # Fix Table 2.1 reference in literature if missing - check table 2 caption
    for p in doc.paragraphs:
        if 'Table 2.1' in p.text and 'Comparison' not in p.text:
            pass

    # Add Table 2.1 caption before comparison table if needed
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('Table 2.1') or (
            i > 0 and 'System Type' in p.text and p.style.name == 'Normal'
        ):
            break

    out = DOC_PATH
    doc.save(str(out))
    print(f'Saved finalized report to {out}')


if __name__ == '__main__':
    process_document()
