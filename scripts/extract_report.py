# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

doc = Document('Nuhu_Muhammad_Datti_Final_Year_Project_Report.docx')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        style = p.style.name if p.style else ''
        print(f'{i:4d}|{style[:25]:25s}|{t[:150]}')

print('\n--- TABLES ---')
for ti, table in enumerate(doc.tables):
    print(f'\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
    for ri, row in enumerate(table.rows[:5]):
        cells = [c.text.strip()[:40] for c in row.cells]
        print(f'  R{ri}: {cells}')

print('\n--- INLINE SHAPES / IMAGES ---')
from docx.document import Document as DocType
# count images via rels
img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
print(f'Images in document: {img_count}')
